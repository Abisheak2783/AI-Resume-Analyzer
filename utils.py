import os
import PyPDF2
import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity


load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    raise ValueError("GEMINI_API_KEY not found. Please configure it in Streamlit Secrets or your .env file.")


genai.configure(api_key=api_key)

model = genai.GenerativeModel("gemini-3.5-flash")
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
def extract_text_from_pdf(pdf_file):
    """Extract text from the uploaded pdf"""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text=""

    for page in pdf_reader.pages:
        extracted = page.extract_text()

        if extracted:
            text+=extracted
    return text

def analyze_resume(resume_text):

    prompt = f"""
You are an experienced HR Recruiter, ATS Specialist, and Career Coach.

Analyze the following resume professionally.

Rules:

1. Never skip any section.
2. If information is unavailable, write "Not Mentioned".
3. Use bullet points.
4. Give ATS Score out of 100.
5. Suggest improvements.
6. Do not make up information.

Resume:

{resume_text}

Provide:

1. Resume Summary

2. Technical Skills

3. Soft Skills

4. Strengths

5. Weaknesses

6. ATS Score

7. Missing Skills

8. Suitable Job Roles

9. Resume Improvement Suggestions
"""

    try:
        response = model.generate_content(prompt)
        return response.text

    except Exception as e:
        return f"Error: {e}"

def compare_resume_job(resume_text, job_description):
    resume = resume_text.lower()
    job = job_description.lower()

    skills = [
        "python",
        "sql",
        "java",
        "c++",
        "django",
        "flask",
        "streamlit",
        "power bi",
        "excel",
        "machine learning",
        "deep learning",
        "git",
        "github",
        "docker",
        "aws",
        "azure",
        "mongodb",
        "mysql",
        "postgresql",
        "html",
        "css",
        "javascript"
    ]

    matched_skills = []
    missing_skills = []

    for skill in skills:
        if skill in resume and skill in job:
            matched_skills.append(skill)

        elif skill in job:
            missing_skills.append(skill)

    total_required = len(matched_skills) + len(missing_skills)

    if total_required == 0:
        score = 0
    else:
        score = round((len(matched_skills) / total_required) * 100)
    return score, matched_skills, missing_skills

#interview questions

def generate_interview_questions(resume_text):

    prompt = f"""
You are a Senior Technical Interviewer with over 15 years of experience.

Analyze the following resume carefully.

Generate interview questions based ONLY on the skills, projects, technologies, and experience mentioned in the resume.

Return the questions in the following format.

## Technical Questions
Generate at least 10 technical questions.

## Project Questions
Generate at least 5 questions based on the projects mentioned.

## HR Questions
Generate at least 5 HR interview questions.

## Coding Questions
Generate at least 5 coding/programming questions.

## Scenario-Based Questions
Generate at least 5 real-world scenario questions.

Do not generate generic questions.
Make the questions personalized to the candidate's resume.

Resume:

{resume_text}
"""

    response = model.generate_content(prompt)

    return response.text

#evaluate answer
def evaluate_answer(question, answer):

    prompt = f"""
You are a Senior Technical Interviewer with 15 years of experience.

Evaluate the candidate's answer.

Question:
{question}

Candidate Answer:
{answer}

Provide your response in the following format:

## Overall Score
Give a score out of 10.

## Technical Accuracy
Evaluate correctness.

## Communication Skills
Evaluate clarity and explanation.

## Strengths
Mention what the candidate did well.

## Weaknesses
Mention mistakes or missing concepts.

## Ideal Answer
Write an ideal interview answer.

## Suggestions
Suggest how the candidate can improve.

Keep the feedback professional.
"""

    response = model.generate_content(prompt)

    return response.text

# ai career advisor
def career_advisor(resume_text):

    prompt = f"""
You are a Senior Career Advisor with 15+ years of experience.

Analyze the following resume carefully.

Based ONLY on the resume, provide a detailed career guidance report.

Return your response in the following format.

# 👨 Candidate Level

Identify whether the candidate is:

- Fresher
- Junior
- Mid-Level
- Senior

Explain why.

--------------------------------------------------

# 🎯 Suitable Job Roles

Suggest the Top 10 suitable job roles.

--------------------------------------------------

# 💰 Expected Salary

Provide:

• India Salary Range

• International Salary Range

--------------------------------------------------

# 📚 Skills to Learn Next

Suggest important skills.

Prioritize them.

--------------------------------------------------

# 📜 Recommended Certifications

Suggest the best certifications.

--------------------------------------------------

# 🚀 Recommended Projects

Suggest 5 portfolio projects.

--------------------------------------------------

# 📅 30-Day Roadmap

Give a learning roadmap.

--------------------------------------------------

# 📅 60-Day Roadmap

Continue the roadmap.

--------------------------------------------------

# 📅 90-Day Roadmap

Continue the roadmap.

--------------------------------------------------

# 🌟 Career Advice

Give professional career advice.

Resume:

{resume_text}

Only provide information related to the resume.
Do not hallucinate.
"""

    response = model.generate_content(prompt)

    return response.text

#Resume Improvement System
def improve_resume(resume_text):

    prompt = f"""
You are an Expert HR Recruiter, ATS Specialist, and Resume Writer.

Analyze the following resume carefully.

Provide a professional resume improvement report.

Return the response in the following format.

# ⭐ Overall Resume Score

Give a score out of 100.

Explain why.

--------------------------------------------------

# 📄 Resume Summary Review

Is the summary good?

How can it be improved?

--------------------------------------------------

# 💻 Technical Skills Review

Mention:

• Existing Skills

• Missing Skills

• Trending Skills

--------------------------------------------------

# 📂 Project Review

Review every project.

Suggest improvements.

Mention if the project description is too short.

--------------------------------------------------

# 📜 Certifications Review

Mention recommended certifications.

--------------------------------------------------

# 🎓 Education Review

Suggest improvements if required.

--------------------------------------------------

# 📈 ATS Optimization

Explain:

• ATS Friendly Score

• Missing Keywords

• Formatting Issues

--------------------------------------------------

# 📝 Resume Writing Suggestions

Suggest:

• Better Action Verbs

• Better Bullet Points

• Better Formatting

--------------------------------------------------

# 🚀 Final Recommendation

Tell whether the resume is:

• Excellent

• Good

• Average

• Needs Improvement

Resume:

{resume_text}

Only analyze the resume.
Do not hallucinate.
"""

    response = model.generate_content(prompt)

    return response.text

#Rewrite Resume
def rewrite_resume(resume_text):

    prompt = f"""
You are a Senior Resume Writer, ATS Specialist, and HR Recruiter.

Rewrite the following resume professionally.

Instructions:

1. Improve the Resume Summary.

2. Improve Technical Skills section.

3. Rewrite every Project professionally.

4. Rewrite Experience professionally.

5. Improve Education section if needed.

6. Improve formatting.

7. Use strong action verbs.

8. Make the resume ATS-Friendly.

9. Keep all information truthful.

10. Do NOT invent experience.

Return the response in the following format.

# Professional Resume

## Resume Summary

...

## Technical Skills

...

## Projects

...

## Experience

...

## Education

...

## Certifications

...

## Final ATS Ready Resume

Resume:

{resume_text}
"""

    response = model.generate_content(prompt)

    return response.text

def semantic_resume_match(resume_text, job_description):

    # Convert resume to embedding
    resume_embedding = embedding_model.encode([resume_text])

    # Convert job description to embedding
    job_embedding = embedding_model.encode([job_description])

    # Calculate cosine similarity
    similarity = cosine_similarity(
        resume_embedding,
        job_embedding
    )[0][0]

    # Convert to percentage
    score = round(similarity * 100, 2)

    return score

#professional resume
def generate_professional_resume(resume_text, template):

    prompt = f"""
You are an expert ATS Resume Writer.

Using the following resume, create a completely professional resume.

Resume:
{resume_text}

The user selected the following template:
{template}

Template Rules:

If ATS Friendly:
- Keep it simple.
- No icons.
- ATS Compatible.
- Professional Summary
- Skills
- Projects
- Experience
- Education
- Certifications

If Modern:
- Attractive formatting.
- Clear headings.
- Professional Summary.
- Skills.
- Projects.
- Experience.
- Education.
- Achievements.

If Corporate:
- Professional business style.
- Executive Summary.
- Skills.
- Experience.
- Projects.
- Education.
- Certifications.

If Minimal:
- Clean.
- Short.
- Only important sections.

Rewrite the resume professionally.

Do NOT create fake experience.

Return only the final resume.
"""

    response = model.generate_content(prompt)

    return response.text

#DOCX Generator Function
def generate_resume_docx(resume_text):

    document = Document()

    document.add_heading("Professional Resume", level=1)

    for line in resume_text.split("\n"):

        document.add_paragraph(line)

    file_name = "Professional_Resume.docx"

    document.save(file_name)

    return file_name

#pdf Generator Function
def generate_resume_pdf(resume_text):

    file_name = "Professional_Resume.pdf"

    doc = SimpleDocTemplate(file_name)

    styles = getSampleStyleSheet()

    story = []

    for line in resume_text.split("\n"):

        story.append(
            Paragraph(line, styles["Normal"])
        )

    doc.build(story)

    return file_name

#Resume Keyword Optimizer
def keyword_optimizer(resume_text, job_description):

    prompt = f"""
You are an ATS Resume Optimization Expert.

Resume:
{resume_text}

Job Description:
{job_description}

Analyze the resume and provide the response in the following format.

1. Important Keywords Found
2. Missing Keywords
3. Recommended Keywords to Add
4. ATS Optimization Tips
5. Final Keyword Optimization Score (out of 100)

Return the response in proper markdown format.
"""

    response = model.generate_content(prompt)

    return response.text

#ats score function
def calculate_ats_score(resume_text):

    score = 100

    suggestions = []

    section_scores = {}

    # Contact Information
    contact = 10
    if "@" not in resume_text:
        contact -= 5
        suggestions.append("Add a professional Email Address.")

    if "linkedin" not in resume_text.lower():
        contact -= 2
        suggestions.append("Add your LinkedIn profile.")

    if "github" not in resume_text.lower():
        contact -= 3
        suggestions.append("Add your GitHub profile.")

    section_scores["Contact Information"] = contact

    # Skills
    skills = 18
    if "python" not in resume_text.lower():
        skills -= 5
        suggestions.append("Mention Python skills.")

    if "sql" not in resume_text.lower():
        skills -= 3
        suggestions.append("Mention SQL.")

    section_scores["Skills"] = skills

    # Projects
    projects = 18
    if "project" not in resume_text.lower():
        projects -= 8
        suggestions.append("Add project experience.")

    section_scores["Projects"] = projects

    # Education
    education = 10
    if "education" not in resume_text.lower():
        education -= 5
        suggestions.append("Add Education section.")

    section_scores["Education"] = education

    # Experience
    experience = 20
    if "experience" not in resume_text.lower():
        experience -= 10
        suggestions.append("Add Work Experience.")

    section_scores["Experience"] = experience

    # Formatting
    formatting = 10
    section_scores["Formatting"] = formatting

    total = sum(section_scores.values())

    return total, section_scores, suggestions



